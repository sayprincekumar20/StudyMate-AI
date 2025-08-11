"""
Streamlit RAG Chatbot (Groq + Chroma + e5 embeddings)
Features:
- Upload documents from UI (pdf/docx/txt)
- Merge new uploads into existing Chroma DB (persistence)
- Keep multiple conversations (chat history)
- Show sources with Download buttons
- Optional reranker (cross-encoder) for better result ordering
"""

import uuid
import time
from datetime import datetime
from typing import List, Dict, Any
import streamlit as st
from pypdf import PdfReader
import docx
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.prompts import PromptTemplate
from groq import Groq
from langchain_core.output_parsers import StrOutputParser


import os
# Load environment variables from .env file
from dotenv import load_dotenv
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    st.session_state.groq_client = None
else:
    st.session_state.groq_client = Groq(api_key=GROQ_API_KEY)

# Optional/-- reranker import will be done lazily if user enables it.
CROSS_ENCODER_MODEL = "cross-encoder/ms-marco-MiniLM-L-6-v2"

# -------------------------
# Config & paths
# -------------------------
PERSIST_DIR = "chroma_store"
UPLOADS_DIR = "uploads"
EMBEDDING_MODEL_NAME = "intfloat/e5-large-v2"  # for retrieval
# Groq model name
GROQ_CHAT_MODEL = "llama-3.1-8b-instant"
# GROQ_CHAT_MODEL = "llama-3.3-70b-versatile"

os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(PERSIST_DIR, exist_ok=True)

# -------------------------
# Streamlit session init
# -------------------------
if "vectordb" not in st.session_state:
    st.session_state.vectordb = None
if "uploaded_files_map" not in st.session_state:
    st.session_state.uploaded_files_map = {}
if "conversations" not in st.session_state:
    st.session_state.conversations = []
if "current_conv" not in st.session_state:
    st.session_state.current_conv = None
if "embedding_model" not in st.session_state:
    st.session_state.embedding_model = None
if "groq_client" not in st.session_state:
    GROQ_API_KEY = os.getenv("GROQ_API_KEY")
    if not GROQ_API_KEY:
        # Streamlit will stop below in UI if no key
        st.session_state.groq_client = None
    else:
        st.session_state.groq_client = Groq(api_key=GROQ_API_KEY)
# lazy reranker
if "reranker" not in st.session_state:
    st.session_state.reranker = None

# -------------------------
# Helpers: file save & loaders
# -------------------------
def save_uploaded_file(u_file) -> Dict[str, Any]:
    """
    Save uploaded file to uploads dir and return metadata.
    Returns dict with keys: key (unique_name), path, orig_name
    """
    orig_name = u_file.name
    unique_id = uuid.uuid4().hex[:8]
    saved_name = f"{unique_id}_{orig_name}"
    saved_path = os.path.join(UPLOADS_DIR, saved_name)
    # uploaded_file may be a SpooledTemporaryFile; use getbuffer if available
    try:
        content = u_file.getbuffer()
    except Exception:
        content = u_file.read()
    with open(saved_path, "wb") as f:
        f.write(content)
    return {"key": saved_name, "path": saved_path, "orig_name": orig_name, "saved_at": time.time()}

def read_text_from_path(path: str) -> str:
    """Load text from a saved file path based on extension."""
    path_lower = path.lower()
    if path_lower.endswith(".txt"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif path_lower.endswith(".pdf"):
        reader = PdfReader(path)
        pages = [p.extract_text() or "" for p in reader.pages]
        return "\n".join(pages)
    elif path_lower.endswith(".docx"):
        doc = docx.Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return ""

# -------------------------
# Embedding model loader (lazy)
# -------------------------
def get_embedding_model():
    if st.session_state.embedding_model is None:
        # This will download / cache the model the first time
        st.session_state.embedding_model = HuggingFaceEmbeddings(
            model_name=EMBEDDING_MODEL_NAME
        )
    return st.session_state.embedding_model

# -------------------------
# Vector DB management
# -------------------------
def load_persisted_vectordb():
    """If a persisted Chroma exists, load it; else None."""
    if os.path.exists(PERSIST_DIR) and os.listdir(PERSIST_DIR):
        emb = get_embedding_model()
        try:
            vectordb = Chroma(persist_directory=PERSIST_DIR, embedding_function=emb)
            return vectordb
        except Exception as e:
            st.warning(f"Failed to load persisted Chroma: {e}")
            return None
    return None


def create_vectordb_from_documents(docs: List[Document], persist: bool = True) -> Chroma:
    """Create a new Chroma from Document objects."""
    emb = get_embedding_model()
    persist_dir = PERSIST_DIR if persist else None

    vectordb = Chroma.from_documents(
        documents=docs,
        embedding=emb,
        persist_directory=persist_dir
    )

    if persist:
        try:
            vectordb._persist()
        except Exception as e:
            print(f"Warning: Failed to persist vector DB: {e}")

    return vectordb


def add_documents_to_vectordb(docs: List[Document]):
    """Add doc chunks to an existing or new vectordb and persist."""
    if not docs:
        return
    if st.session_state.vectordb is None:
        st.session_state.vectordb = create_vectordb_from_documents(docs, persist=True)
    else:
        st.session_state.vectordb.add_documents(docs)
        try:
            st.session_state.vectordb.persist()
        except Exception:
            pass

# -------------------------
# Chunking helper
# -------------------------
def chunk_documents(file_docs: List[Dict[str, Any]], chunk_size=1000, chunk_overlap=200) -> List[Document]:
    """
    file_docs: list of dicts with keys path, orig_name, key
    Returns List[Document] for embedding.
    """
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    chunks = []
    for f in file_docs:
        text = read_text_from_path(f["path"])
        if not text or text.strip() == "":
            continue
        raw_chunks = splitter.split_text(text)
        for idx, c in enumerate(raw_chunks):
            metadata = {"source_key": f["key"], "filename": f["orig_name"], "chunk": idx, "path": f["path"]}
            chunks.append(Document(page_content=c, metadata=metadata))
    return chunks

# -------------------------
# Reranker lazy init
# -------------------------
def init_reranker():
    if st.session_state.reranker is None:
        # lazy import to avoid heavy startup cost if not used
        from sentence_transformers import CrossEncoder
        st.session_state.reranker = CrossEncoder(CROSS_ENCODER_MODEL)
    return st.session_state.reranker

def rerank_candidates(query: str, docs: List[Document], top_k=5) -> List[Document]:
    reranker = init_reranker()
    texts = [d.page_content for d in docs]
    pairs = [[query, t] for t in texts]
    scores = reranker.predict(pairs)
    order = list(reversed(sorted(range(len(scores)), key=lambda i: scores[i])))
    ordered = [docs[i] for i in order]
    return ordered[:top_k]

# -------------------------
# Groq chat call and generation
# -------------------------
from typing import Tuple

def generate_answer_with_groq(query: str, top_k_retrieve=10, rerank=False, top_k_final=5) -> Tuple[str, List[Dict[str, str]]]:
    if st.session_state.vectordb is None:
        return "No documents in KB. Please upload documents first.", []

    retriever = st.session_state.vectordb.as_retriever(search_kwargs={"k": top_k_retrieve})
    raw_docs = retriever.invoke(query)
    # Optional rerank
    if rerank and len(raw_docs) > top_k_final:
        docs = rerank_candidates(query, raw_docs, top_k=top_k_final)
    else:
        docs = raw_docs[:top_k_final]

    # Build context with citation metadata
    context_parts = []
    seen_sources = {}
    for d in docs:
        fn = d.metadata.get("filename", d.metadata.get("source_key", "unknown"))
        src_key = d.metadata.get("source_key", None)
        snippet = d.page_content.strip()
        if len(snippet) > 1500:
            snippet = snippet[:1500] + "..."
        context_parts.append(f"Source: {fn}\n\n{snippet}\n")
        # record path for download buttons (de-duplicate)
        if src_key and src_key not in seen_sources:
            # find saved path from session_state.uploaded_files_map
            fileinfo = st.session_state.uploaded_files_map.get(src_key)
            if fileinfo:
                seen_sources[src_key] = {"filename": fileinfo["orig_name"], "path": fileinfo["path"]}

    context_text = "\n\n---\n\n".join(context_parts) if context_parts else "No context found."

    # Prompt template
    prompt_template = """
You are a helpful assistant. Use ONLY the following context to answer the question when possible.
If the answer is not present in the context, say "I couldn't find an answer in the provided materials."

CONTEXT:
{context}

QUESTION:
{question}

INSTRUCTIONS:
- Provide a short accurate answer.
- Add a SOURCES section listing the filenames used (one per line).
"""
    prompt = PromptTemplate.from_template(prompt_template)
    filled = prompt.format(context=context_text, question=query)

    # Call Groq (chat completion)
    groq_client = st.session_state.groq_client
    if not groq_client:
        return "Groq API key not configured. Set GROQ_API_KEY env variable.", []

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_CHAT_MODEL,
            messages=[{"role": "user", "content": filled}],
            temperature=0.0,
        )
        # path into response depends on SDK; using .choices[0].message.content like earlier examples
        raw_text = completion.choices[0].message.content
    except Exception as e:
        return f"Error calling Groq: {e}", []

    parser = StrOutputParser()
    parsed = parser.parse(raw_text)
    # return parsed answer and list of sources (filename + path)
    sources_list = list(seen_sources.values())
    return parsed, sources_list

# -------------------------
# Conversation utilities
# -------------------------
def new_conversation(title: str = None): # pyright: ignore[reportArgumentType]
    conv_id = uuid.uuid4().hex[:8]
    title = title or f"Conversation {len(st.session_state.conversations) + 1}"
    conv = {"id": conv_id, "title": title, "history": [], "created_at": time.time()}
    st.session_state.conversations.insert(0, conv)  # newest first
    st.session_state.current_conv = conv_id
    return conv_id


def get_current_conversation():
    conv_id = st.session_state.current_conv
    if conv_id is None:
        # create default one
        conv_id = new_conversation("Welcome Conversation")
    for c in st.session_state.conversations:
        if c["id"] == conv_id:
            return c
    # fallback
    return new_conversation()

def save_message_to_current_conv(question: str, answer: str, sources: List[Dict[str,str]]):
    conv = get_current_conversation()
    entry = {
        "question": question,
        "answer": answer,
        "sources": sources,
        "ts": time.time()
    }
    if isinstance(conv, dict) and "history" in conv:
        conv["history"].append(entry)

# -------------------------
# UI Layout (Streamlit)
# -------------------------
st.set_page_config(page_title="StudyMate AI — Learning Assistant", layout="wide")
# Main layout: left sidebar, main area
sidebar = st.sidebar
sidebar.title("StudyMate AI")
sidebar.markdown("Your intelligent learning companion")

# Upload section in sidebar (allows merging)
sidebar.subheader("Upload Documents")
uploaded_files = sidebar.file_uploader(
    "Add PDFs / DOCX / TXT (multiple)",
    type=["pdf", "docx", "txt"],
    accept_multiple_files=True
)
if uploaded_files:
    if sidebar.button("Add to Knowledge Base"):
        # Save uploaded files locally and chunk+embed them
        saved_meta = []
        for uf in uploaded_files:
            info = save_uploaded_file(uf)
            st.session_state.uploaded_files_map[info["key"]] = info
            saved_meta.append(info)
        # chunk and add to vectordb
        chunks = chunk_documents(saved_meta)
        if not chunks:
            st.warning("No text extracted from uploaded files.")
        else:
            add_documents_to_vectordb(chunks)
            sidebar.success(f"Added {len(saved_meta)} files ({len(chunks)} chunks) to KB.")
            # If no current conversation exists, create one
            if st.session_state.current_conv is None:
                new_conversation()

# list recent conversations (click to switch)
sidebar.markdown("### Recent Conversations")
for conv in st.session_state.conversations:
    # show a compact button; when pressed, set current_conv
    if sidebar.button(conv["title"], key=f"conv_{conv['id']}"):
        st.session_state.current_conv = conv["id"]

# New conversation button
if sidebar.button("➕ New Conversation"):
    new_conversation()

# Option: load existing persisted Chroma if present
if st.session_state.vectordb is None:
    st.session_state.vectordb = load_persisted_vectordb()
    if st.session_state.vectordb:
        sidebar.info("Loaded persisted knowledge base.")

# Optional settings
sidebar.markdown("---")
enable_rerank = sidebar.checkbox("Enable re-ranker (better ordering, slower)", value=False)
top_k_retrieve = sidebar.number_input("Retriever: top k", min_value=1, max_value=50, value=10, step=1)
top_k_final = sidebar.number_input("Final top k to send to LLM", min_value=1, max_value=10, value=5, step=1)

# Main area header

# 1️ HEADER (Light shaded background)
sidebar_bg_color = "#f0f2f6"  

st.markdown(
    f"""
    <div style='background-color: {sidebar_bg_color}; padding: 10px; border-radius: 0px; 
                border-bottom: 1px solid #ddd;'>
        <h3 style='margin: 0;'>Learning Assistant</h3>
        <p style='margin: 0;'>Ask me anything about your studies</p>
    </div>
    """,
    unsafe_allow_html=True
)
# Small space
st.markdown("")

# 2️ IMAGE (Hero/Banner)
# Centered Image with Fixed Width
import base64
hero_path = "/mnt/data/Screenshot 2025-08-10 004136.png"
try:
    if os.path.exists(hero_path):
        with open(hero_path, "rb") as img_file:
            img_bytes = img_file.read()
            img_base64 = base64.b64encode(img_bytes).decode("utf-8")
        st.markdown(
            f"""
           <div style="text-align: center; margin-top: 10px;">
                <img src="data:image/png;base64,{img_base64}" 
                     style="width: 400px; max-width: 100%; border-radius: 15px;">
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            """
             <div style="text-align: center; margin-top: 10px;">
                <img src="https://images.unsplash.com/photo-1522202176988-66273c2fd55f?q=80&w=1200&auto=format&fit=crop&ixlib=rb-4.0.3&s=placeholder"
                     style="width: 400px; max-width: 100%; border-radius: 15px;">
            </div>
            """,
            unsafe_allow_html=True
        )
except Exception:
    pass


# 3️ SHORT DESCRIPTION (3 lines)
st.markdown(
    """
    <div style='text-align: center; font-size: 16px; line-height: 1.5; margin-top: 15px;'>
        <b>Welcome to <span style="color:#7B4EFF;">StudyMate AI</span></b><br>
        Your intelligent learning companion powered by advanced AI.<br>
        Upload your study materials, ask questions, and get personalized explanations<br>
        tailored to your learning style and pace.
    </div>
    """,
    unsafe_allow_html=True
)

st.markdown("---")  # Separator before rest of UI


# Display conversation history
conv = get_current_conversation()
if isinstance(conv, dict) and 'title' in conv:
    st.markdown(f"## {conv['title']}")
else:
    st.markdown("## Conversation")

history_area = st.container()
with history_area:
    if not (isinstance(conv, dict) and "history" in conv and conv["history"]): 
        st.info("No messages yet. Upload documents and ask a question to start.")
    else:
        # Show each QA pair
        for entry in conv["history"]: 
            q_ts = datetime.fromtimestamp(entry["ts"]).strftime("%Y-%m-%d %H:%M:%S") 
            st.markdown(f"**You** ({q_ts}):")
            st.write(entry["question"]) 
            st.markdown("**StudyMate AI:**")
            st.write(entry["answer"]) 
            if entry["sources"]:
                st.markdown("**Sources:**")
                for idx, s in enumerate(entry["sources"]):
                    fname = s["filename"]
                    fpath = s["path"]
                    # Download button for each source
                    try:
                        with open(fpath, "rb") as f:
                            file_bytes = f.read()
                        st.download_button(label=f"Download: {fname}", data=file_bytes, file_name=fname, key=f"dl_{entry['ts']}_{idx}")
                    except Exception as e:
                        st.write(f"- {fname} (file not found)")

st.markdown("---")

# Input area (bottom)
user_question = st.text_input("Ask me anything about your study materials...", key="user_input")
send = st.button("Send")

if send and user_question and user_question.strip():
    if st.session_state.vectordb is None:
        st.warning("No documents in KB. Please upload documents first.")
    else:
        with st.spinner("Searching and generating answer..."):
            # ensure reranker is initialized if enabled
            if enable_rerank:
                init_reranker()
            answer, sources = generate_answer_with_groq(user_question, top_k_retrieve, rerank=enable_rerank, top_k_final=top_k_final)
        # Save to current conversation
        save_message_to_current_conv(user_question, answer, sources)
        # optionally rename the conversation title to first question snippet if it is default
        if isinstance(conv, dict) and "title" in conv and (conv["title"].startswith("Conversation") or conv["title"].startswith("Welcome")):
            conv["title"] = user_question[:40] + ("..." if len(user_question) > 40 else "")
        # re-run to display updated history (Streamlit auto-reruns after this block)
        st.rerun()
()
